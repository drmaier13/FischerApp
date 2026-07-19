from __future__ import annotations

import json
import sys
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


def main() -> None:
    path = Path(sys.argv[1])
    document = Document(path)

    paragraph_rows = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph_rows.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": bool(run.underline),
                        "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                        "highlight": str(run.font.highlight_color) if run.font.highlight_color else None,
                    }
                    for run in paragraph.runs
                    if run.text
                ],
            }
        )

    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        comments = []
        if "word/comments.xml" in names:
            root = ET.fromstring(archive.read("word/comments.xml"))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for node in root.findall("w:comment", ns):
                comments.append(
                    {
                        "id": node.attrib.get(f"{{{ns['w']}}}id"),
                        "author": node.attrib.get(f"{{{ns['w']}}}author"),
                        "text": "".join(node.itertext()).strip(),
                    }
                )
        media = [name for name in names if name.startswith("word/media/")]

    summary = {
        "path": str(path),
        "paragraph_count": len(document.paragraphs),
        "non_empty_paragraph_count": len(paragraph_rows),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "styles": Counter(row["style"] for row in paragraph_rows),
        "comments": comments,
        "media": media,
        "sample_paragraphs": paragraph_rows[:180],
        "sample_tables": [
            [[cell.text for cell in row.cells] for row in table.rows[:10]]
            for table in document.tables[:5]
        ],
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2, default=lambda value: dict(value)))


if __name__ == "__main__":
    main()
